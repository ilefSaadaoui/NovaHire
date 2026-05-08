"""
Pipeline Stage: Multi-Format Text Extraction
Responsibilities:
  1. Detect input format (PDF / DOCX / image)
  2. PDF  → pdfplumber (layout-rich) → PyMuPDF (fallback)
  3. DOCX → docx2txt (texte + liens) → python-docx (fallback)
  4. Image / scanned PDF → pytesseract (principal FR/EN) → EasyOCR (fallback complexe)
  5. Conditional OCR only when text is below threshold (lazy-loaded)

Anti-perte : 3 méthodes en cascade pour chaque format, cache SHA-256.
This module never calls scoring, parsing, or any other stage.
"""

import io
import logging
import re
import shutil
from pathlib import Path
from typing import Optional

import pdfplumber
import torch

from config import settings
from models import ExtractionResult
from cache import extraction_cache, content_hash

logger = logging.getLogger("NovaHire-AI.extractor")

# ── Vérification Tesseract au démarrage ─────────────────────
# Sur Windows, le PATH système n'est pas toujours chargé par le processus Python.
# On configure le chemin explicitement via pytesseract avant toute détection.
try:
    import pytesseract as _pyt
    import os as _os

    # Priorité 1 : chemin explicite depuis config
    if settings.tesseract_cmd and _os.path.isfile(settings.tesseract_cmd):
        _pyt.pytesseract.tesseract_cmd = settings.tesseract_cmd
        _TESSERACT_AVAILABLE = True
        logger.info(f"Tesseract configure: {settings.tesseract_cmd}")
    else:
        # Priorité 2 : chercher dans le PATH système
        _TESSERACT_AVAILABLE = shutil.which("tesseract") is not None
        if _TESSERACT_AVAILABLE:
            logger.info("Tesseract detecte dans le PATH — OCR principal active.")
        else:
            logger.warning(
                "Tesseract non trouve. OCR utilisera EasyOCR en fallback. "
                "Installer depuis: https://github.com/UB-Mannheim/tesseract/releases"
            )
except ImportError:
    _TESSERACT_AVAILABLE = False
    logger.warning("pytesseract non installe — OCR Tesseract desactive.")



def _get_easyocr_reader():
    """Lazy-load EasyOCR uniquement en dernier recours."""
    global _easyocr_reader, _easyocr_load_attempted
    if _easyocr_load_attempted:
        return _easyocr_reader
    _easyocr_load_attempted = True
    try:
        import easyocr
        logger.info("Lazy-loading EasyOCR (fallback OCR — cas complexe détecté)...")
        _easyocr_reader = easyocr.Reader(
            settings.ocr_languages, gpu=torch.cuda.is_available()
        )
        logger.info("EasyOCR chargé avec succès (fallback prêt).")
    except Exception as e:
        logger.warning(f"EasyOCR non disponible: {e}")
        _easyocr_reader = None
    return _easyocr_reader


# ══════════════════════════════════════════════
# Détection du format
# ══════════════════════════════════════════════

def _detect_format(filename: str, file_content: bytes) -> str:
    """
    Détecte le format par extension puis par magic bytes.
    Retourne: 'pdf' | 'docx' | 'image'
    """
    ext = Path(filename).suffix.lower().lstrip(".")
    if ext == "pdf":
        return "pdf"
    if ext in ("docx", "doc"):
        return "docx"
    if ext in ("png", "jpg", "jpeg", "tiff", "bmp"):
        return "image"

    # Magic bytes fallback (sans extension fiable)
    if file_content[:4] == b"%PDF":
        return "pdf"
    if file_content[:2] == b"PK":  # ZIP = DOCX/XLSX
        return "docx"
    return "image"


# ══════════════════════════════════════════════
# Extracteurs PDF
# ══════════════════════════════════════════════

def _extract_pdfplumber(file_content: bytes) -> tuple[str, int]:
    """Extraction via pdfplumber. Meilleur pour CV à mise en page complexe."""
    text = ""
    page_count = 0
    try:
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            page_count = len(pdf.pages)
            for page in pdf.pages:
                text += page.extract_text() or ""
    except Exception as e:
        logger.warning(f"pdfplumber échoué: {e}")
    return text.strip(), page_count


def _extract_pymupdf(file_content: bytes) -> tuple[str, int]:
    """Fallback PDF via PyMuPDF (rapide, robuste)."""
    text = ""
    page_count = 0
    try:
        import fitz
        with fitz.open(stream=file_content, filetype="pdf") as doc:
            page_count = len(doc)
            for page in doc:
                text += page.get_text()
    except Exception as e:
        logger.error(f"PyMuPDF échoué: {e}")
    return text.strip(), page_count


# ══════════════════════════════════════════════
# Extracteurs DOCX
# ══════════════════════════════════════════════

def _extract_docx2txt(file_content: bytes) -> str:
    """
    Extraction DOCX via docx2txt.
    Avantage : extrait texte + hyperlinks + texte des images alt.
    """
    try:
        import docx2txt
        text = docx2txt.process(io.BytesIO(file_content))
        return text.strip() if text else ""
    except Exception as e:
        logger.warning(f"docx2txt échoué: {e}")
        return ""


def _extract_python_docx(file_content: bytes) -> str:
    """
    Fallback DOCX via python-docx.
    Extrait paragraphes + texte des tableaux.
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # Texte dans les tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"python-docx échoué: {e}")
        return ""


# ══════════════════════════════════════════════
# Extracteurs OCR (images / PDF scannés)
# ══════════════════════════════════════════════

def _render_pdf_to_images(file_content: bytes) -> list:
    """Rend les pages d'un PDF en images PIL pour Tesseract."""
    try:
        import fitz
        from PIL import Image as PILImage
        doc = fitz.open(stream=file_content, filetype="pdf")
        images = []
        scale = settings.ocr_dpi_scale
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
            img = PILImage.frombytes("RGB", [pix.width, pix.height], pix.samples)
            images.append(img)
        return images
    except Exception as e:
        logger.error(f"Rendu PDF→image échoué: {e}")
        return []


def _extract_tesseract(file_content: bytes, fmt: str) -> str:
    """
    OCR principal via pytesseract (Tesseract).
    Plus rapide qu'EasyOCR, support natif FR+EN, ne charge pas de modèle lourd.
    Retourne '' si Tesseract n'est pas installé (fallback EasyOCR sera utilisé).
    """
    if not _TESSERACT_AVAILABLE:
        return ""
    try:
        import pytesseract
        from PIL import Image as PILImage

        config = f"--oem 3 --psm 6 -l {settings.tesseract_lang}"

        if fmt == "image":
            img = PILImage.open(io.BytesIO(file_content))
            return pytesseract.image_to_string(img, config=config).strip()

        # PDF scanné → rendu page par page
        images = _render_pdf_to_images(file_content)
        if not images:
            return ""
        parts = [pytesseract.image_to_string(img, config=config) for img in images]
        return "\n".join(parts).strip()

    except ImportError:
        logger.warning("pytesseract non installé — passage à EasyOCR.")
        return ""
    except Exception as e:
        logger.error(f"Tesseract OCR échoué: {e}")
        return ""


def _extract_easyocr_fallback(file_content: bytes, fmt: str) -> str:
    """
    OCR fallback via EasyOCR (lazy-loaded, ~500 MB RAM).
    Utilisé uniquement si Tesseract échoue ou retourne vide.
    """
    reader = _get_easyocr_reader()
    if reader is None:
        return ""
    try:
        import fitz
        from PIL import Image as PILImage

        if fmt == "image":
            from PIL import Image as PILImage
            img = PILImage.open(io.BytesIO(file_content))
            results = reader.readtext(img, detail=0)
            return " ".join(results).strip()

        doc = fitz.open(stream=file_content, filetype="pdf")
        ocr_text = ""
        scale = settings.ocr_dpi_scale
        for page in doc:
            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
            img_data = pix.tobytes("png")
            results = reader.readtext(img_data, detail=0)
            ocr_text += " ".join(results) + "\n"
        return ocr_text.strip()
    except Exception as e:
        logger.error(f"EasyOCR fallback échoué: {e}")
        return ""


# ══════════════════════════════════════════════
# Point d'entrée principal
# ══════════════════════════════════════════════

def extract_text(file_content: bytes, filename: str = "resume.pdf") -> ExtractionResult:
    """
    Extraction multi-format avec pipeline anti-perte :

    PDF  → pdfplumber → PyMuPDF → (OCR si texte < seuil)
    DOCX → docx2txt   → python-docx
    Image/Scanné → Tesseract → EasyOCR

    Résultat mis en cache par hash SHA-256 pour idempotence.
    """
    # ── Cache check ──────────────────────────
    file_hash = content_hash(file_content)
    cached = extraction_cache.get(file_hash)
    if cached is not None:
        logger.info(f"Cache extraction HIT ({file_hash[:12]}...)")
        return cached

    fmt = _detect_format(filename, file_content)
    logger.info(f"Format détecté: '{fmt}' pour '{filename}'")

    text = ""
    method = "unknown"
    page_count = 0

    # ── Branche PDF ──────────────────────────
    if fmt == "pdf":
        text, page_count = _extract_pdfplumber(file_content)
        method = "pdfplumber"

        if not text:
            logger.info("pdfplumber vide → fallback PyMuPDF...")
            text, page_count = _extract_pymupdf(file_content)
            method = "pymupdf"

        # OCR conditionnel si texte insuffisant (PDF scanné)
        if len(text) < settings.ocr_min_text_threshold:
            logger.info(
                f"Texte insuffisant ({len(text)} chars) → OCR Tesseract..."
            )
            ocr_text = _extract_tesseract(file_content, "pdf")
            if ocr_text:
                text = ocr_text
                method = "tesseract-pdf"
            else:
                logger.info("Tesseract vide → EasyOCR fallback...")
                ocr_text = _extract_easyocr_fallback(file_content, "pdf")
                if ocr_text:
                    text = ocr_text
                    method = "easyocr-pdf"

    # ── Branche DOCX ─────────────────────────
    elif fmt == "docx":
        text = _extract_docx2txt(file_content)
        method = "docx2txt"

        if not text:
            logger.info("docx2txt vide → fallback python-docx...")
            text = _extract_python_docx(file_content)
            method = "python-docx"

        page_count = 1  # DOCX n'a pas de notion de page fixe

    # ── Branche Image ─────────────────────────
    else:
        logger.info("Format image → OCR Tesseract...")
        text = _extract_tesseract(file_content, "image")
        method = "tesseract-image"

        if not text:
            logger.info("Tesseract image vide → EasyOCR fallback...")
            text = _extract_easyocr_fallback(file_content, "image")
            method = "easyocr-image"
        page_count = 1

    # ── Nettoyage des artefacts PDF ───────────
    text = _clean_extracted_text(text)

    # ── Résultat + cache ─────────────────────
    result = ExtractionResult(
        raw_text=text,
        method=method,
        file_format=fmt,
        page_count=page_count,
        char_count=len(text),
    )

    extraction_cache.put(file_hash, result)
    logger.info(
        f"Extraction terminée: {result.char_count} chars via [{method}] "
        f"({page_count} page(s), format={fmt})"
    )
    return result


# ══════════════════════════════════════════════
# Nettoyage des artefacts d'encodage PDF
# ══════════════════════════════════════════════

_CID_PATTERN = re.compile(r'\(cid:\d+\)', re.IGNORECASE)
_REPEATED_CHAR_PATTERN = re.compile(r'(.)\1{2,}')  # 3+ identical chars in a row

# Characters that CAN legitimately repeat (dashes, dots, spaces, underscores)
_ALLOWED_REPEAT_CHARS = set('-–—_.• =*#~')


def _clean_extracted_text(text: str) -> str:
    """
    Remove common PDF font-encoding artifacts from extracted text:

    1. (cid:127), (cid:32), etc.  — unresolved character IDs from embedded
       fonts that pdfplumber cannot decode.  These appear when the PDF uses
       a custom glyph map (e.g. Wingdings bullets, custom separators).

    2. 'nnnnn', 'aaaaa', 'fffff', etc. — runs of 3+ identical alpha characters
       that are clearly garbled glyphs (real words rarely repeat a letter 3+
       times in a row).  We keep repeated punctuation / dashes unchanged
       because those can be legitimate (e.g. '---' as a separator).
    """
    if not text:
        return text

    # Step 1: strip (cid:NNN) tokens entirely
    text = _CID_PATTERN.sub('', text)

    # Step 2: collapse runs of identical *alphabetic* characters
    def _replace_run(m: re.Match) -> str:
        char = m.group(1)
        if char in _ALLOWED_REPEAT_CHARS:
            return m.group(0)   # keep legitimate repeats
        if char.isalpha():
            return ''           # remove garbled glyph runs
        return m.group(0)       # keep numeric / other repeats

    text = _REPEATED_CHAR_PATTERN.sub(_replace_run, text)

    # Step 3: collapse multiple consecutive spaces / clean up leftover whitespace
    text = re.sub(r'[ \t]{2,}', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()
